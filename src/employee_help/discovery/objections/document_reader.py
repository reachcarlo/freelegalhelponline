"""Document reader for the objection drafter — extracts text from .docx and .pdf uploads."""

from __future__ import annotations

import structlog

logger = structlog.get_logger(__name__)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
ALLOWED_EXTENSIONS = {".docx", ".pdf"}


class DocumentReadError(Exception):
    """Raised when a document cannot be read."""


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from an uploaded file.

    Args:
        file_bytes: Raw file contents.
        filename: Original filename (used for extension detection).

    Returns:
        Extracted text content.

    Raises:
        DocumentReadError: On unsupported type, oversize, or extraction failure.
    """
    if len(file_bytes) > MAX_FILE_SIZE:
        raise DocumentReadError(
            f"File too large ({len(file_bytes) / 1024 / 1024:.1f} MB). "
            f"Maximum allowed size is {MAX_FILE_SIZE // 1024 // 1024} MB."
        )

    ext = _get_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise DocumentReadError(
            f"Unsupported file type: {ext}. Only .docx and .pdf files are accepted."
        )

    if ext == ".docx":
        return extract_text_from_docx(file_bytes)
    return extract_text_from_pdf(file_bytes)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file using python-docx.

    Extracts paragraph text and table cell text.
    """
    import io

    from docx import Document

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise DocumentReadError(f"Failed to read .docx file: {exc}") from exc

    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    result = "\n".join(parts)
    if not result.strip():
        raise DocumentReadError("Document appears to be empty — no text content found.")
    return result


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a .pdf file using pdfplumber."""
    import io

    import pdfplumber

    try:
        pdf = pdfplumber.open(io.BytesIO(file_bytes))
    except Exception as exc:
        raise DocumentReadError(f"Failed to read .pdf file: {exc}") from exc

    parts: list[str] = []
    try:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text.strip())
    finally:
        pdf.close()

    result = "\n".join(parts)
    if not result.strip():
        raise DocumentReadError("PDF appears to be empty — no text content found.")
    return result


def _get_extension(filename: str) -> str:
    """Extract lowercase file extension."""
    dot = filename.rfind(".")
    if dot == -1:
        return ""
    return filename[dot:].lower()
