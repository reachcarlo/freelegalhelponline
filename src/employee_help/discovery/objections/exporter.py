"""Word document exporter for the objection drafter.

Generates standalone .docx files and inserts objections into response shell documents.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from employee_help.discovery.objections.models import DISCLAIMER, WAIVER_PREAMBLE

logger = structlog.get_logger(__name__)

# Matches "RESPONSE TO ... NO. X:" markers — same pattern used in parser.py
_RESPONSE_MARKER = re.compile(
    r"^\s*RESPONSE\s+TO\s+(?:SPECIAL\s+)?(?:INTERROGATOR(?:Y|IES)|REQUEST|DEMAND|RFA|RFP)"
    r"\s+(?:FOR\s+(?:PRODUCTION|ADMISSION)\s+(?:OF\s+DOCUMENTS?\s+)?)?"
    r"(?:NO\.?\s*)?(\d+)\s*[:.)\s]",
    re.IGNORECASE,
)

# Detects italic-worthy case names: "Name v. Name"
_CASE_NAME_PATTERN = re.compile(r"([A-Z][A-Za-z\s.,]+\s+v\.\s+[A-Z][A-Za-z\s.,]+)")


@dataclass
class ExportOptions:
    """Options controlling .docx export content."""

    include_request_text: bool = False
    include_waiver_language: bool = False
    enabled_objections: dict[str, bool] = field(default_factory=dict)


def generate_standalone_docx(
    results: list[dict],
    options: ExportOptions | None = None,
) -> bytes:
    """Generate a standalone .docx with objection responses.

    Args:
        results: List of RequestAnalysisInfo dicts (from API response).
        options: Export options controlling content inclusion.

    Returns:
        .docx file bytes.
    """
    opts = options or ExportOptions()
    doc = Document()

    # Page setup: 1" margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default font: 12pt Times New Roman
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    for result in results:
        request_number = result["request_number"]
        request_text = result.get("request_text", "")
        discovery_type = result.get("discovery_type", "interrogatories")
        objections = result.get("objections", [])
        no_objections_rationale = result.get("no_objections_rationale")

        # Filter to enabled objections
        enabled = [
            o for o in objections
            if opts.enabled_objections.get(
                f"{request_number}-{o['ground_id']}", True
            )
        ]

        type_label = _discovery_type_singular(discovery_type).upper()

        # Request text (optional)
        if opts.include_request_text:
            p = doc.add_paragraph()
            run = p.add_run(f"{type_label} NO. {request_number}:")
            run.bold = True
            _set_font(run, "Times New Roman", Pt(12))
            doc.add_paragraph(request_text)

        # Response header
        p = doc.add_paragraph()
        run = p.add_run(f"RESPONSE TO {type_label} NO. {request_number}:")
        run.bold = True
        _set_font(run, "Times New Roman", Pt(12))

        if not enabled and no_objections_rationale:
            doc.add_paragraph(no_objections_rationale)
        else:
            for obj in enabled:
                _add_objection_paragraph(doc, obj)

        # Waiver preamble
        if opts.include_waiver_language and enabled:
            doc.add_paragraph(WAIVER_PREAMBLE)

        # Blank line between requests
        doc.add_paragraph()

    # Disclaimer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(DISCLAIMER)
    run.italic = True
    _set_font(run, "Times New Roman", Pt(10))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def insert_into_shell(
    shell_bytes: bytes,
    results: list[dict],
    options: ExportOptions | None = None,
) -> tuple[bytes, int, int]:
    """Insert objections into an uploaded response shell .docx.

    Finds "RESPONSE TO ... NO. X:" markers and inserts objection paragraphs
    after each matched marker.

    Args:
        shell_bytes: Raw .docx bytes of the response shell.
        results: List of RequestAnalysisInfo dicts.
        options: Export options.

    Returns:
        Tuple of (modified_docx_bytes, markers_filled, total_markers).
    """
    opts = options or ExportOptions()
    doc = Document(io.BytesIO(shell_bytes))

    # Index results by request number
    results_by_number: dict[int, dict] = {}
    for r in results:
        num = r["request_number"]
        if isinstance(num, int):
            results_by_number[num] = r

    markers_filled = 0
    total_markers = 0

    # Iterate paragraphs and find RESPONSE TO markers
    for para in list(doc.paragraphs):
        m = _RESPONSE_MARKER.match(para.text)
        if not m:
            continue

        total_markers += 1
        request_num = int(m.group(1))
        result = results_by_number.get(request_num)
        if not result:
            continue

        objections = result.get("objections", [])
        enabled = [
            o for o in objections
            if opts.enabled_objections.get(
                f"{request_num}-{o['ground_id']}", True
            )
        ]

        if not enabled:
            no_rationale = result.get("no_objections_rationale")
            if no_rationale:
                _insert_paragraph_after(para, no_rationale)
                markers_filled += 1
            continue

        # Insert objection paragraphs after the marker, in reverse order
        # so they end up in the correct order
        if opts.include_waiver_language:
            _insert_paragraph_after(para, WAIVER_PREAMBLE)

        for obj in reversed(enabled):
            text = _build_objection_text(obj)
            _insert_paragraph_after(para, text)

        markers_filled += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), markers_filled, total_markers


def _add_objection_paragraph(doc: Document, obj: dict) -> None:
    """Add a single objection as a paragraph with italic case names."""
    label = obj.get("label", "")
    explanation = obj.get("explanation", "")
    statutory = obj.get("statutory_citations", [])
    cases = obj.get("case_citations", [])

    p = doc.add_paragraph()

    # "Objection — {label}: "
    run = p.add_run(f"Objection — {label}: ")
    run.bold = True
    _set_font(run, "Times New Roman", Pt(12))

    # Explanation
    run = p.add_run(explanation)
    _set_font(run, "Times New Roman", Pt(12))

    # Citations
    cite_parts: list[str] = []
    if statutory:
        stat_text = ", ".join(f"{c['code']} {c['section']}" for c in statutory)
        cite_parts.append(stat_text)

    if cases:
        for c in cases:
            cite_parts.append(f"{c['name']} {c['citation']}")

    if cite_parts:
        run = p.add_run(" (")
        _set_font(run, "Times New Roman", Pt(12))

        # Add case citations with italic names
        for i, c in enumerate(cases):
            if i > 0 or statutory:
                run = p.add_run("; ")
                _set_font(run, "Times New Roman", Pt(12))

            if i == 0 and statutory:
                # Add statutory first
                stat_text = ", ".join(
                    f"{s['code']} {s['section']}" for s in statutory
                )
                run = p.add_run(stat_text + "; ")
                _set_font(run, "Times New Roman", Pt(12))

            # Italic case name
            run = p.add_run(c["name"])
            run.italic = True
            _set_font(run, "Times New Roman", Pt(12))

            run = p.add_run(f" {c['citation']}")
            _set_font(run, "Times New Roman", Pt(12))

        if not cases and statutory:
            stat_text = ", ".join(
                f"{s['code']} {s['section']}" for s in statutory
            )
            run = p.add_run(stat_text)
            _set_font(run, "Times New Roman", Pt(12))

        run = p.add_run(")")
        _set_font(run, "Times New Roman", Pt(12))


def _build_objection_text(obj: dict) -> str:
    """Build plain text for a single objection (for shell insertion)."""
    label = obj.get("label", "")
    explanation = obj.get("explanation", "")
    statutory = obj.get("statutory_citations", [])
    cases = obj.get("case_citations", [])

    cite_parts: list[str] = []
    if statutory:
        cite_parts.append(
            ", ".join(f"{c['code']} {c['section']}" for c in statutory)
        )
    if cases:
        for c in cases:
            cite_parts.append(f"{c['name']} {c['citation']}")

    cites = "; ".join(cite_parts)
    text = f"Objection — {label}: {explanation}"
    if cites:
        text += f" ({cites})"
    return text


def _insert_paragraph_after(para, text: str) -> None:
    """Insert a new paragraph with text after the given paragraph element."""
    new_p = para._element.makeelement(qn("w:p"), {})
    new_r = new_p.makeelement(qn("w:r"), {})
    new_t = new_r.makeelement(qn("w:t"), {})
    new_t.text = text
    # Preserve leading/trailing spaces
    new_t.set(qn("xml:space"), "preserve")
    new_r.append(new_t)
    new_p.append(new_r)
    para._element.addnext(new_p)


def _set_font(run, name: str, size) -> None:
    """Set font name and size on a Run."""
    run.font.name = name
    run.font.size = size


def _discovery_type_singular(dtype: str) -> str:
    """Convert discovery type value to singular label."""
    mapping = {
        "interrogatories": "Interrogatory",
        "rfps": "Request for Production",
        "rfas": "Request for Admission",
    }
    return mapping.get(dtype, "Request")
