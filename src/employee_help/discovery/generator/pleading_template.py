"""California 28-line pleading paper DOCX template generator.

Creates a Word document template that conforms to California Rules
of Court, Rules 2.100–2.111 for pleading paper:

  - 28 numbered lines per page
  - Line numbers in left margin
  - 1" top/bottom margins, 1" left margin (after line numbers), 0.5" right margin
  - 12pt Courier New (monospaced) or proportional font per CRC 2.105
  - Case caption on first page
  - Page numbers in footer

The generated template uses Jinja2 tags (via docxtpl) so it can be
filled with case data and discovery content.

Pure computation — no DB, no ML, no external services.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, Twips

# ---------------------------------------------------------------------------
# Constants per California Rules of Court 2.100–2.111
# ---------------------------------------------------------------------------

LINES_PER_PAGE = 28
FONT_NAME = "Times New Roman"  # CRC 2.105 allows proportional
FONT_SIZE_PT = 12
LINE_SPACING_PT = 24  # Double-spaced = 24pt for 12pt font

# Margins
TOP_MARGIN = Inches(1.0)
BOTTOM_MARGIN = Inches(0.75)
LEFT_MARGIN = Inches(1.5)  # Includes space for line numbers
RIGHT_MARGIN = Inches(0.5)

# Line number column width
LINE_NUMBER_GUTTER = Inches(0.5)


def _add_line_numbering(section) -> None:
    """Add line numbering to a Word document section.

    CRC 2.108 requires consecutive line numbers on each page, starting
    from 1, in the left margin.
    """
    sectPr = section._sectPr
    ln_num = OxmlElement("w:lnNumType")
    ln_num.set(qn("w:countBy"), "1")
    ln_num.set(qn("w:start"), "1")
    ln_num.set(qn("w:restart"), "newPage")
    # Distance from text: about 0.3 inches
    ln_num.set(qn("w:distance"), str(Twips(216)))  # 216 twips = ~0.15"
    sectPr.append(ln_num)


def _add_page_number_footer(section) -> None:
    """Add centered page number to the section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # PAGE field
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run3._element.append(fldChar2)

    run.font.name = FONT_NAME
    run.font.size = Pt(10)


def _set_paragraph_format(paragraph, font_size: int = FONT_SIZE_PT) -> None:
    """Apply standard pleading paper formatting to a paragraph."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(LINE_SPACING_PT)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(font_size)


def create_pleading_template(
    output_path: Path | str,
    *,
    include_caption: bool = True,
) -> Path:
    """Create a blank California pleading paper DOCX template with Jinja2 tags.

    The template includes:
    - Proper page layout (margins, orientation)
    - Line numbering in left margin
    - Page numbers in footer
    - Jinja2 template tags for docxtpl filling

    Args:
        output_path: Where to save the generated .docx file.
        include_caption: Whether to include the case caption block.

    Returns:
        The output path.
    """
    doc = Document()
    output_path = Path(output_path)

    # ----- Page setup -----
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = TOP_MARGIN
    section.bottom_margin = BOTTOM_MARGIN
    section.left_margin = LEFT_MARGIN
    section.right_margin = RIGHT_MARGIN

    _add_line_numbering(section)
    _add_page_number_footer(section)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_PT)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # Clear any default paragraphs
    for p in doc.paragraphs:
        p.clear()

    if include_caption:
        _add_caption_block(doc)

    # ----- Document title (Jinja2 tag) -----
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("{{ document_title }}")
    title_run.bold = True
    title_run.font.name = FONT_NAME
    title_run.font.size = Pt(FONT_SIZE_PT)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(title_p)

    # Blank line after title
    doc.add_paragraph()

    # ----- Propounding/Responding party line -----
    party_p = doc.add_paragraph()
    party_run = party_p.add_run(
        "Propounding Party: {{ propounding_party }}    "
        "Responding Party: {{ responding_party }}    "
        "Set No.: {{ set_number }}"
    )
    party_run.font.name = FONT_NAME
    party_run.font.size = Pt(10)
    _set_paragraph_format(party_p, font_size=10)

    # Blank line
    doc.add_paragraph()

    # ----- Definitions section (conditional) -----
    def_p = doc.add_paragraph()
    def_run = def_p.add_run("{% if definitions %}")
    def_run.font.name = FONT_NAME
    def_run.font.size = Pt(FONT_SIZE_PT)
    _set_paragraph_format(def_p)

    def_title = doc.add_paragraph()
    def_title_run = def_title.add_run("DEFINITIONS")
    def_title_run.bold = True
    def_title_run.font.name = FONT_NAME
    def_title_run.font.size = Pt(FONT_SIZE_PT)
    def_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(def_title)

    doc.add_paragraph()

    # Loop through definitions
    def_loop = doc.add_paragraph()
    def_loop_run = def_loop.add_run("{% for term, definition in definitions.items() %}")
    def_loop_run.font.name = FONT_NAME
    def_loop_run.font.size = Pt(FONT_SIZE_PT)
    _set_paragraph_format(def_loop)

    def_item = doc.add_paragraph()
    def_item_run = def_item.add_run(
        '     "{{ term }}" means {{ definition }}'
    )
    def_item_run.font.name = FONT_NAME
    def_item_run.font.size = Pt(FONT_SIZE_PT)
    _set_paragraph_format(def_item)

    doc.add_paragraph()

    def_endloop = doc.add_paragraph()
    def_endloop.add_run("{% endfor %}")
    _set_paragraph_format(def_endloop)

    def_endif = doc.add_paragraph()
    def_endif.add_run("{% endif %}")
    _set_paragraph_format(def_endif)

    doc.add_paragraph()

    # ----- Production instructions (RFPD only, conditional) -----
    pi_if = doc.add_paragraph()
    pi_if.add_run("{% if production_instructions %}")
    _set_paragraph_format(pi_if)

    pi_title = doc.add_paragraph()
    pi_title_run = pi_title.add_run("PRELIMINARY STATEMENT AND PRODUCTION INSTRUCTIONS")
    pi_title_run.bold = True
    pi_title_run.font.name = FONT_NAME
    pi_title_run.font.size = Pt(FONT_SIZE_PT)
    pi_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(pi_title)

    doc.add_paragraph()

    pi_text = doc.add_paragraph()
    pi_text.add_run("{{ production_instructions }}")
    _set_paragraph_format(pi_text)

    doc.add_paragraph()

    pi_endif = doc.add_paragraph()
    pi_endif.add_run("{% endif %}")
    _set_paragraph_format(pi_endif)

    # ----- Requests section -----
    req_title = doc.add_paragraph()
    req_title_run = req_title.add_run("{{ requests_heading }}")
    req_title_run.bold = True
    req_title_run.font.name = FONT_NAME
    req_title_run.font.size = Pt(FONT_SIZE_PT)
    req_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(req_title)

    doc.add_paragraph()

    # Loop through requests
    req_loop = doc.add_paragraph()
    req_loop.add_run("{% for request in requests %}")
    _set_paragraph_format(req_loop)

    req_label = doc.add_paragraph()
    req_label_run = req_label.add_run("     {{ request.label }}")
    req_label_run.bold = True
    req_label_run.font.name = FONT_NAME
    req_label_run.font.size = Pt(FONT_SIZE_PT)
    _set_paragraph_format(req_label)

    doc.add_paragraph()

    req_text = doc.add_paragraph()
    req_text.add_run("     {{ request.text }}")
    _set_paragraph_format(req_text)

    doc.add_paragraph()

    req_endloop = doc.add_paragraph()
    req_endloop.add_run("{% endfor %}")
    _set_paragraph_format(req_endloop)

    # ----- Signature block -----
    doc.add_paragraph()

    sig_date = doc.add_paragraph()
    sig_date.add_run("Dated: {{ date }}")
    _set_paragraph_format(sig_date)

    doc.add_paragraph()
    doc.add_paragraph()

    sig_line = doc.add_paragraph()
    sig_line.add_run("                                        ____________________________________")
    _set_paragraph_format(sig_line)

    sig_name = doc.add_paragraph()
    sig_name.add_run("                                        {{ attorney_name }}")
    _set_paragraph_format(sig_name)

    sig_for = doc.add_paragraph()
    sig_for.add_run("                                        Attorney for {{ attorney_for }}")
    _set_paragraph_format(sig_for)

    # ----- Declaration of Necessity (conditional, SROGs/RFAs > 35) -----
    _add_declaration_section(doc)

    doc.save(str(output_path))
    return output_path


def _add_declaration_section(doc: Document) -> None:
    """Add a conditional Declaration of Necessity page.

    Rendered only when the ``declaration`` context dict is provided
    (i.e., when SROGs > 35 or fact RFAs > 35).

    CCP 2030.050 (SROGs) / CCP 2033.050 (RFAs).
    """
    # Conditional open
    decl_if = doc.add_paragraph()
    decl_if.add_run("{% if declaration %}")
    _set_paragraph_format(decl_if)

    # Page break
    pb = doc.add_paragraph()
    run_pb = pb.add_run()
    run_pb.add_break(WD_BREAK.PAGE)
    _set_paragraph_format(pb)

    # Heading
    heading = doc.add_paragraph()
    heading_run = heading.add_run("DECLARATION FOR ADDITIONAL DISCOVERY")
    heading_run.bold = True
    heading_run.font.name = FONT_NAME
    heading_run.font.size = Pt(FONT_SIZE_PT)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(heading)

    # CCP section subheading
    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        "(Code of Civil Procedure \u00a7 {{ declaration.ccp_section }})"
    )
    sub_run.font.name = FONT_NAME
    sub_run.font.size = Pt(FONT_SIZE_PT)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(sub)

    doc.add_paragraph()

    # Declarant intro
    intro = doc.add_paragraph()
    intro.add_run("I, {{ declaration.declarant_name }}, declare:")
    _set_paragraph_format(intro)

    doc.add_paragraph()

    # Numbered paragraphs
    _paras = [
        (
            "1.  I am the attorney for {{ declaration.attorney_for }} "
            "in this action."
        ),
        (
            "2.  This discovery set contains {{ declaration.request_count }} "
            "{{ declaration.request_type_plural }}, which exceeds the limit of "
            "35 set forth in CCP \u00a7 {{ declaration.limit_section }}."
        ),
        (
            "3.  I have personally examined each "
            "{{ declaration.request_type_singular }} in this set."
        ),
        (
            "4.  None of the {{ declaration.request_type_plural }} in this set "
            "is being propounded for any improper purpose, such as to harass, "
            "cause unnecessary delay, or needlessly increase the cost of litigation."
        ),
        (
            "5.  None of the {{ declaration.request_type_plural }} is "
            "unreasonably cumulative or duplicative, or can be obtained from "
            "some other source that is more convenient, less burdensome, or "
            "less expensive."
        ),
        (
            "6.  Each {{ declaration.request_type_singular }} in this set is "
            "warranted because {{ declaration.justification }}."
        ),
    ]
    for text in _paras:
        p = doc.add_paragraph()
        p.add_run(text)
        _set_paragraph_format(p)
        doc.add_paragraph()

    # Penalty of perjury
    perjury = doc.add_paragraph()
    perjury.add_run(
        "I declare under penalty of perjury under the laws of the State of "
        "California that the foregoing is true and correct."
    )
    _set_paragraph_format(perjury)

    doc.add_paragraph()

    # Date
    decl_date = doc.add_paragraph()
    decl_date.add_run("Dated: {{ date }}")
    _set_paragraph_format(decl_date)

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature line
    decl_sig = doc.add_paragraph()
    decl_sig.add_run(
        "                                        ____________________________________"
    )
    _set_paragraph_format(decl_sig)

    decl_name = doc.add_paragraph()
    decl_name.add_run(
        "                                        {{ declaration.declarant_name }}"
    )
    _set_paragraph_format(decl_name)

    # Conditional close
    decl_endif = doc.add_paragraph()
    decl_endif.add_run("{% endif %}")
    _set_paragraph_format(decl_endif)


def _add_caption_block(doc: Document) -> None:
    """Add the case caption block to the first page.

    Standard California Superior Court caption with:
    - Attorney info (upper left)
    - Court name and county
    - Party names with "vs."
    - Case number
    """
    # Attorney header block
    for tag in [
        "{{ attorney_name }}{% if attorney_sbn %} (SBN {{ attorney_sbn }}){% endif %}",
        "{% if attorney_firm %}{{ attorney_firm }}{% endif %}",
        "{{ attorney_address }}",
        "{{ attorney_city_state_zip }}",
        "Telephone: {{ attorney_phone }}",
        "{% if attorney_fax %}Facsimile: {{ attorney_fax }}{% endif %}",
        "Email: {{ attorney_email }}",
        "",
        "Attorney for {{ attorney_for }}",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(tag)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(12)  # Single-spaced for header
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    # Blank line before court name
    doc.add_paragraph()

    # Court name
    court_p = doc.add_paragraph()
    court_run = court_p.add_run(
        "SUPERIOR COURT OF THE STATE OF CALIFORNIA"
    )
    court_run.bold = True
    court_run.font.name = FONT_NAME
    court_run.font.size = Pt(FONT_SIZE_PT)
    court_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(court_p)

    county_p = doc.add_paragraph()
    county_run = county_p.add_run(
        "COUNTY OF {{ court_county | upper }}"
    )
    county_run.bold = True
    county_run.font.name = FONT_NAME
    county_run.font.size = Pt(FONT_SIZE_PT)
    county_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(county_p)

    doc.add_paragraph()

    # Plaintiff block
    plt_p = doc.add_paragraph()
    plt_p.add_run("{{ plaintiff_block }},")
    _set_paragraph_format(plt_p)

    # Indented "Plaintiff(s),"
    desig_p1 = doc.add_paragraph()
    desig_p1.add_run("                    Plaintiff(s),")
    _set_paragraph_format(desig_p1)

    # "vs." line with case number to the right
    vs_p = doc.add_paragraph()
    vs_p.add_run(
        "     vs."
        "                                        Case No. {{ case_number }}"
    )
    _set_paragraph_format(vs_p)

    doc.add_paragraph()

    # Defendant block
    def_p = doc.add_paragraph()
    def_p.add_run("{{ defendant_block }},")
    _set_paragraph_format(def_p)

    desig_p2 = doc.add_paragraph()
    desig_p2.add_run("                    Defendant(s).")
    _set_paragraph_format(desig_p2)

    # Horizontal rule (underscores)
    rule_p = doc.add_paragraph()
    rule_p.add_run("_" * 60)
    _set_paragraph_format(rule_p)

    doc.add_paragraph()
