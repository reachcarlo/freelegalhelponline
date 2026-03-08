"""Proof of Service DOCX generator for California discovery.

Generates a standard California Proof of Service document (equivalent to
POS-030) as a Word document using python-docx + docxtpl on the same
pleading paper template used for SROGs/RFPDs/RFAs.

Accepted by California courts in lieu of the Judicial Council POS-030
form when the information content is equivalent.

Pure computation — no DB, no ML, no external services.
"""

from __future__ import annotations

import tempfile
from datetime import date
from pathlib import Path

from ..case_info import (
    defendant_block,
    plaintiff_block,
    propounding_party_name,
)
from ..models import CaseInfo, ServiceMethod
from .pleading_template import create_pleading_template

# ---------------------------------------------------------------------------
# Template directory
# ---------------------------------------------------------------------------

TEMPLATE_DIR = Path(__file__).parent / "templates"

# ---------------------------------------------------------------------------
# Service method descriptions per CCP 1013 / 1010.6
# ---------------------------------------------------------------------------

SERVICE_METHOD_DESCRIPTIONS: dict[ServiceMethod, str] = {
    ServiceMethod.PERSONAL: (
        "by personally delivering a true and correct copy thereof to the "
        "person(s) listed below"
    ),
    ServiceMethod.MAIL_IN_STATE: (
        "by placing a true and correct copy thereof enclosed in a sealed "
        "envelope with postage thereon fully prepaid, in the United States "
        "mail at the address set forth below (CCP § 1013(a))"
    ),
    ServiceMethod.MAIL_OUT_OF_STATE: (
        "by placing a true and correct copy thereof enclosed in a sealed "
        "envelope with postage thereon fully prepaid, in the United States "
        "mail to an out-of-state address set forth below (CCP § 1013(a))"
    ),
    ServiceMethod.MAIL_INTERNATIONAL: (
        "by placing a true and correct copy thereof enclosed in a sealed "
        "envelope with postage thereon fully prepaid, in the United States "
        "mail to an international address set forth below (CCP § 1013(a))"
    ),
    ServiceMethod.ELECTRONIC: (
        "by transmitting a true and correct copy thereof by electronic "
        "service to the person(s) at the electronic service address(es) "
        "set forth below (CCP § 1010.6)"
    ),
    ServiceMethod.OVERNIGHT: (
        "by placing a true and correct copy thereof enclosed in a sealed "
        "envelope for overnight delivery to the address set forth below "
        "(CCP § 1013(c))"
    ),
}


def _get_or_create_template() -> Path:
    """Get the pleading paper template, creating it if needed."""
    template_path = TEMPLATE_DIR / "pos_template.docx"
    if not template_path.exists():
        _create_pos_template(template_path)
    return template_path


def _create_pos_template(output_path: Path) -> Path:
    """Create a POS-specific pleading paper template with Jinja2 tags.

    Reuses the caption block from pleading_template but has POS-specific
    body content instead of discovery requests.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Pt

    from .pleading_template import (
        FONT_NAME,
        FONT_SIZE_PT,
        LINE_SPACING_PT,
        _add_caption_block,
        _add_line_numbering,
        _add_page_number_footer,
        _set_paragraph_format,
        BOTTOM_MARGIN,
        LEFT_MARGIN,
        RIGHT_MARGIN,
        TOP_MARGIN,
    )
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches

    doc = Document()
    output_path = Path(output_path)

    # Page setup
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = TOP_MARGIN
    section.bottom_margin = BOTTOM_MARGIN
    section.left_margin = LEFT_MARGIN
    section.right_margin = RIGHT_MARGIN

    _add_line_numbering(section)
    _add_page_number_footer(section)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_PT)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    for p in doc.paragraphs:
        p.clear()

    # Caption
    _add_caption_block(doc)

    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("PROOF OF SERVICE")
    title_run.bold = True
    title_run.font.name = FONT_NAME
    title_run.font.size = Pt(FONT_SIZE_PT)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(title_p)

    doc.add_paragraph()

    # Declarant intro
    intro = doc.add_paragraph()
    intro.add_run("I, {{ server_name }}, declare:")
    _set_paragraph_format(intro)

    doc.add_paragraph()

    # Paragraph 1: Employment/residence
    p1 = doc.add_paragraph()
    p1.add_run(
        "1.  At the time of service I was over 18 years of age and not a "
        "party to this action. My business address is {{ server_address }}."
    )
    _set_paragraph_format(p1)

    doc.add_paragraph()

    # Paragraph 2: Service date and method
    p2 = doc.add_paragraph()
    p2.add_run(
        "2.  On {{ service_date }}, I served the following document(s):"
    )
    _set_paragraph_format(p2)

    doc.add_paragraph()

    # Document list
    doc_loop = doc.add_paragraph()
    doc_loop.add_run("{% for doc_name in documents_served %}")
    _set_paragraph_format(doc_loop)

    doc_item = doc.add_paragraph()
    doc_item.add_run("     \u2022  {{ doc_name }}")
    _set_paragraph_format(doc_item)

    doc_endloop = doc.add_paragraph()
    doc_endloop.add_run("{% endfor %}")
    _set_paragraph_format(doc_endloop)

    doc.add_paragraph()

    # Paragraph 3: Service method
    p3 = doc.add_paragraph()
    p3.add_run("3.  {{ service_method_description }}")
    _set_paragraph_format(p3)

    doc.add_paragraph()

    # Paragraph 4: Person served
    p4 = doc.add_paragraph()
    p4.add_run(
        "4.  The document(s) were served on the following party:"
    )
    _set_paragraph_format(p4)

    doc.add_paragraph()

    served = doc.add_paragraph()
    served.add_run(
        "     {{ served_party_name }}\n"
        "     {{ served_party_address }}"
    )
    _set_paragraph_format(served)

    doc.add_paragraph()

    # Penalty of perjury
    perjury = doc.add_paragraph()
    perjury.add_run(
        "I declare under penalty of perjury under the laws of the State of "
        "California that the foregoing is true and correct."
    )
    _set_paragraph_format(perjury)

    doc.add_paragraph()

    # Date + signature
    sig_date = doc.add_paragraph()
    sig_date.add_run("Dated: {{ service_date }}")
    _set_paragraph_format(sig_date)

    doc.add_paragraph()
    doc.add_paragraph()

    sig_line = doc.add_paragraph()
    sig_line.add_run(
        "                                        ____________________________________"
    )
    _set_paragraph_format(sig_line)

    sig_name = doc.add_paragraph()
    sig_name.add_run(
        "                                        {{ server_name }}"
    )
    _set_paragraph_format(sig_name)

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def build_proof_of_service(
    case_info: CaseInfo,
    *,
    server_name: str,
    server_address: str,
    served_party_name: str,
    served_party_address: str,
    service_method: ServiceMethod,
    service_date: date,
    documents_served: list[str],
) -> bytes:
    """Build a Proof of Service DOCX for discovery documents.

    Args:
        case_info: Case and attorney information for the caption.
        server_name: Name of the person who performed service.
        server_address: Business address of the server.
        served_party_name: Name of the party being served.
        served_party_address: Address where documents were served.
        service_method: How the documents were served.
        service_date: Date of service.
        documents_served: List of document names served.

    Returns:
        Filled DOCX as bytes.
    """
    method_desc = SERVICE_METHOD_DESCRIPTIONS[service_method]

    context = {
        # Caption
        "attorney_name": case_info.attorney.name,
        "attorney_sbn": case_info.attorney.sbn,
        "attorney_firm": case_info.attorney.firm_name or "",
        "attorney_address": case_info.attorney.address,
        "attorney_city_state_zip": case_info.attorney.city_state_zip,
        "attorney_phone": case_info.attorney.phone,
        "attorney_fax": case_info.attorney.fax or "",
        "attorney_email": case_info.attorney.email,
        "attorney_for": (
            case_info.attorney.attorney_for
            or propounding_party_name(case_info)
        ),
        "court_county": case_info.court_county,
        "plaintiff_block": plaintiff_block(case_info),
        "defendant_block": defendant_block(case_info),
        "case_number": case_info.case_number,
        # POS content
        "server_name": server_name,
        "server_address": server_address,
        "served_party_name": served_party_name,
        "served_party_address": served_party_address,
        "service_method_description": method_desc,
        "service_date": service_date.strftime("%B %d, %Y"),
        "documents_served": documents_served,
    }

    from docxtpl import DocxTemplate

    template_path = _get_or_create_template()
    tpl = DocxTemplate(str(template_path))
    tpl.render(context)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tpl.save(tmp.name)
        tmp.seek(0)
        return Path(tmp.name).read_bytes()
