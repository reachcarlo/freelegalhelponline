"""Tests for LITIGAGENT DocxExtractor (L1.5)."""

from __future__ import annotations

import io

import pytest

from docx import Document  # type: ignore[import-untyped]

from employee_help.casefile.extractors.base import ExtractionResult, FileExtractor
from employee_help.casefile.extractors.docx import DocxExtractor


# --- Helpers ---


def _make_docx(
    paragraphs: list[str] | None = None,
    headings: list[tuple[str, int]] | None = None,
    tables: list[list[list[str]]] | None = None,
    header_text: str | None = None,
    footer_text: str | None = None,
) -> bytes:
    """Build a minimal .docx in memory and return its bytes.

    Args:
        paragraphs: List of plain paragraph texts.
        headings: List of (text, level) tuples.
        tables: List of tables, each a list of rows (list of cell strings).
        header_text: Text to place in the document header.
        footer_text: Text to place in the document footer.
    """
    doc = Document()

    if header_text:
        section = doc.sections[0]
        section.header.is_linked_to_previous = False
        section.header.paragraphs[0].text = header_text

    if footer_text:
        section = doc.sections[0]
        section.footer.is_linked_to_previous = False
        section.footer.paragraphs[0].text = footer_text

    if headings:
        for text, level in headings:
            doc.add_heading(text, level=level)

    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    if tables:
        for table_data in tables:
            if not table_data:
                continue
            num_cols = len(table_data[0])
            tbl = doc.add_table(rows=len(table_data), cols=num_cols)
            for r, row_data in enumerate(table_data):
                for c, cell_text in enumerate(row_data):
                    tbl.cell(r, c).text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- Interface tests ---


class TestDocxExtractorInterface:
    def test_can_extract_by_extension(self):
        assert DocxExtractor().can_extract("application/octet-stream", "docx") is True

    def test_can_extract_by_mime(self):
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert DocxExtractor().can_extract(mime, "unknown") is True

    def test_cannot_extract_other(self):
        ext = DocxExtractor()
        assert ext.can_extract("application/pdf", "pdf") is False
        assert ext.can_extract("text/plain", "txt") is False
        assert ext.can_extract("application/msword", "doc") is False

    def test_supported_extensions(self):
        assert DocxExtractor().supported_extensions == {"docx"}

    def test_isinstance_file_extractor(self):
        assert isinstance(DocxExtractor(), FileExtractor)


# --- Paragraph extraction ---


class TestDocxParagraphExtraction:
    def test_single_paragraph(self):
        data = _make_docx(paragraphs=["Hello world"])
        result = DocxExtractor().extract(data, "test.docx")

        assert "Hello world" in result.text
        assert isinstance(result, ExtractionResult)
        assert result.metadata["extractor"] == "docx"
        assert result.warnings == []

    def test_multiple_paragraphs(self):
        paras = ["First paragraph.", "Second paragraph.", "Third paragraph."]
        data = _make_docx(paragraphs=paras)
        result = DocxExtractor().extract(data, "multi.docx")

        for p in paras:
            assert p in result.text

    def test_empty_document(self):
        data = _make_docx()
        result = DocxExtractor().extract(data, "empty.docx")

        assert result.text.strip() == ""
        assert any("No text extracted" in w for w in result.warnings)


# --- Heading extraction ---


class TestDocxHeadingExtraction:
    def test_heading_preserved_as_markdown(self):
        data = _make_docx(headings=[("Main Title", 1)])
        result = DocxExtractor().extract(data, "heading.docx")

        assert "# Main Title" in result.text

    def test_heading_levels(self):
        data = _make_docx(headings=[
            ("Level 1", 1),
            ("Level 2", 2),
            ("Level 3", 3),
        ])
        result = DocxExtractor().extract(data, "levels.docx")

        assert "# Level 1" in result.text
        assert "## Level 2" in result.text
        assert "### Level 3" in result.text

    def test_headings_and_paragraphs_mixed(self):
        data = _make_docx(
            headings=[("Introduction", 1)],
            paragraphs=["This is the intro text."],
        )
        result = DocxExtractor().extract(data, "mixed.docx")

        assert "# Introduction" in result.text
        assert "This is the intro text." in result.text


# --- Table extraction ---


class TestDocxTableExtraction:
    def test_simple_table(self):
        table_data = [
            ["Name", "Age"],
            ["Alice", "30"],
            ["Bob", "25"],
        ]
        data = _make_docx(tables=[table_data])
        result = DocxExtractor().extract(data, "table.docx")

        assert "Name" in result.text
        assert "Alice" in result.text
        assert "Bob" in result.text
        # Markdown table separators
        assert "---" in result.text
        assert "|" in result.text

    def test_table_pipe_escaping(self):
        table_data = [["Header"], ["Value with | pipe"]]
        data = _make_docx(tables=[table_data])
        result = DocxExtractor().extract(data, "pipe.docx")

        assert "\\|" in result.text

    def test_multiple_tables(self):
        t1 = [["A", "B"], ["1", "2"]]
        t2 = [["X", "Y"], ["3", "4"]]
        data = _make_docx(tables=[t1, t2])
        result = DocxExtractor().extract(data, "multi_table.docx")

        assert "A" in result.text
        assert "X" in result.text


# --- Header/footer extraction ---


class TestDocxHeaderFooter:
    def test_header_extracted(self):
        data = _make_docx(header_text="CONFIDENTIAL", paragraphs=["Body text."])
        result = DocxExtractor().extract(data, "header.docx")

        assert "CONFIDENTIAL" in result.text
        assert "Body text." in result.text

    def test_footer_extracted(self):
        data = _make_docx(footer_text="Page 1 of 1", paragraphs=["Body text."])
        result = DocxExtractor().extract(data, "footer.docx")

        assert "Page 1 of 1" in result.text

    def test_header_and_footer(self):
        data = _make_docx(
            header_text="HEADER",
            footer_text="FOOTER",
            paragraphs=["Middle."],
        )
        result = DocxExtractor().extract(data, "hf.docx")

        assert "HEADER" in result.text
        assert "FOOTER" in result.text
        assert "Middle." in result.text


# --- Error handling ---


class TestDocxErrorHandling:
    def test_invalid_bytes(self):
        result = DocxExtractor().extract(b"not a docx file", "bad.docx")

        assert result.text == ""
        assert any("Failed to open" in w for w in result.warnings)
        assert result.metadata["extractor"] == "docx"

    def test_returns_extraction_result(self):
        data = _make_docx(paragraphs=["Test"])
        result = DocxExtractor().extract(data, "test.docx")
        assert isinstance(result, ExtractionResult)
